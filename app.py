import streamlit as st
import google.generativeai as genai
import os
import PyPDF2 as pdf
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import base64
import unicodedata
import re

# Load environment variables
load_dotenv()

# Configure Gemini API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Gemini Pro Response Function
def get_gemini_response(input_text):
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(input_text)
    return response.text

# PDF Reader Function
def input_pdf_text(uploaded_file):
    reader = pdf.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += str(page.extract_text())
    return text

# Prompt Template
def build_prompt(resume_text, jd_text):
    return f"""
You are an advanced ATS (Application Tracking System) with expertise in matching tech resumes to job descriptions.
Analyze the resume and JD, and provide:
1. Match percentage
2. List of missing keywords
3. Suggested profile summary to enhance alignment.

Resume:
{resume_text}

Job Description:
{jd_text}

Respond in JSON format with:
{{"JD Match":"%","MissingKeywords":[],"Profile Summary":""}}
"""

# Resume Update Prompt
def build_resume_update_prompt(resume_text, jd_text):
    return f"""
You are a professional resume writer. Based on the given job description and existing resume text, update the resume to improve its match. Ensure it retains a professional tone and incorporates missing keywords from the JD.
Also, rewrite or add a Projects section with 2–3 relevant project titles and descriptions tailored to the job.
Maintain formatting with bold headings, bullet points (•), and horizontal rules (like "—") separating sections.
Include a heading at the top with the candidate's name and contact info.
Include a LinkedIn and GitHub section if applicable.

Existing Resume:
{resume_text}

Job Description:
{jd_text}

Return the improved resume in plain text format.
"""

# Enhanced Styled Resume PDF Generator
def save_as_pdf(text, filename="updated_resume.pdf"):
    pdf_writer = FPDF()
    pdf_writer.add_page()
    pdf_writer.set_auto_page_break(auto=True, margin=15)
    pdf_writer.set_font("Arial", size=11)

    # Add header with contact info
    pdf_writer.set_font("Arial", style="B", size=14)
    pdf_writer.cell(0, 10, "PISKA ROHITH", ln=True)
    pdf_writer.set_font("Arial", size=11)
    pdf_writer.cell(0, 8, "Hyderabad - rohithpiska32@gmail.com - +91-9347466130", ln=True)
    pdf_writer.cell(0, 8, "linkedin.com/in/PiskaRohith - github.com/rohithnani143", ln=True)
    pdf_writer.ln(4)
    pdf_writer.cell(190, 0, "", ln=True, border="T")
    pdf_writer.ln(5)

    # Normalize and sanitize text
    text = text.replace("•", "-")
    text = unicodedata.normalize("NFKD", text).encode("latin-1", "ignore").decode("latin-1")
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if line == "":
            pdf_writer.ln(4)
        elif re.match(r"^[A-Z ]{3,}$", line):
            pdf_writer.set_font("Arial", style="B", size=12)
            pdf_writer.cell(0, 10, line, ln=True)
            pdf_writer.set_font("Arial", size=11)
        elif line.startswith("-"):
            pdf_writer.cell(5)
            pdf_writer.multi_cell(0, 8, line)
        elif re.match(r"^[-–—]{5,}$", line):
            pdf_writer.cell(0, 5, "", ln=True)
            pdf_writer.cell(190, 0, "", ln=True, border="T")
            pdf_writer.cell(0, 5, "", ln=True)
        else:
            pdf_writer.multi_cell(0, 8, line)

    path = os.path.join(tempfile.gettempdir(), filename)
    pdf_writer.output(path)
    return path

# Save as DOCX

def save_as_docx(text, filename="updated_resume.docx"):
    doc = Document()

    # Header
    header = doc.add_heading("PISKA ROHITH", level=1)
    para = doc.add_paragraph("Hyderabad • rohithpiska32@gmail.com • +91-9347466130\n")
    para.add_run("linkedin.com/in/PiskaRohith • github.com/rohithnani143")
    doc.add_paragraph("\n")

    for line in text.split('\n'):
        if re.match(r"^[A-Z ]{3,}$", line):
            doc.add_heading(line, level=2)
        elif line.startswith("•"):
            doc.add_paragraph(line, style='List Bullet')
        elif re.match(r"^[-–—]{5,}$", line):
            doc.add_paragraph("\n––––––––––––––––––––––––––––––––––––––––––\n")
        else:
            doc.add_paragraph(line)

    path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(path)
    return path

# Download link generator
def generate_download_link(file_path, label, ext):
    with open(file_path, "rb") as f:
        base64_data = base64.b64encode(f.read()).decode('utf-8')
    return f'<a href="data:application/{ext};base64,{base64_data}" download="{os.path.basename(file_path)}">{label}</a>'

# Streamlit Page Config
st.set_page_config(page_title="Smart ATS AI", layout="wide", page_icon="📄")

# Header
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    h1, .stTextInput, .stButton>button, .stFileUploader>div {
        font-family: 'Segoe UI', sans-serif;
    }
    .css-1aumxhk { background-color: #ffffff; border-radius: 10px; padding: 20px; }
    </style>
""", unsafe_allow_html=True)

st.title("📄 Smart ATS Resume Analyzer & Enhancer")
st.markdown("Improve your resume with AI insights and download an optimized version!")

# Input Columns
col1, col2 = st.columns(2)
with col1:
    jd = st.text_area("📌 Paste Job Description", height=300, placeholder="Enter the JD here...")
with col2:
    uploaded_file = st.file_uploader("📁 Upload Your Resume (PDF)", type="pdf", help="Only PDF files are supported")

if st.button("🚀 Analyze & Improve Resume"):
    if uploaded_file and jd:
        resume_text = input_pdf_text(uploaded_file)

        # Analysis
        analysis_prompt = build_prompt(resume_text, jd)
        with st.spinner("Analyzing resume..."):
            response = get_gemini_response(analysis_prompt)
        st.subheader("✅ ATS Feedback")
        st.code(response, language="json")

        # Resume update
        update_prompt = build_resume_update_prompt(resume_text, jd)
        with st.spinner("Enhancing resume based on analysis..."):
            updated_resume = get_gemini_response(update_prompt)
        st.subheader("📝 Updated Resume Preview")
        st.text_area("Updated Resume", updated_resume, height=400)

        # Save and download PDF + DOCX
        pdf_path = save_as_pdf(updated_resume)
        docx_path = save_as_docx(updated_resume)
        st.markdown(generate_download_link(pdf_path, "📄 Download PDF", "pdf"), unsafe_allow_html=True)
        st.markdown(generate_download_link(docx_path, "📝 Download Word Document", "vnd.openxmlformats-officedocument.wordprocessingml.document"), unsafe_allow_html=True)

    else:
        st.warning("Please upload a resume and paste the job description to begin analysis.")

# Footer
st.markdown("---")
st.markdown("Made with ❤️ by Rohith | Powered by Gemini")
